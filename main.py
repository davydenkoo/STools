from datetime import datetime, timedelta
import configparser
import json
import os
from pathlib import Path, PurePosixPath
from time import sleep
import re

from docx import Document
import ftputil
from PIL import Image
import requests


parser = configparser.ConfigParser()

parser.read("config.ini", encoding="utf-8")

FORECAST_ROOT_FOLDER = parser.get("FORECAST", "FORECAST_ROOT_FOLDER")

FORECAST_IMG_SRC_FOLDER = parser.get("FORECAST", "FORECAST_IMG_SRC_FOLDER")
FORECAST_IMG_DST_FOLDER = parser.get("FORECAST", "FORECAST_IMG_DST_FOLDER")

FULL_IMG_SIZE = (1920, 520)
FB_IMG_SIZE = (1200, 630)
INTRO_IMG_SIZE = (478, 264)

FULL_IMG_TAIL = "_f.jpg"
FB_IMG_TAIL = "_og_tw.jpg"
INTRO_IMG_TAIL = "_i.jpg"

FULL_IMG_FILL_LAYER_COLOR = (17, 28, 36)

# FULL_IMG_FILL_LAYER_OPACITY = 80
FULL_IMG_FILL_LAYER_OPACITY = 70
# FULL_IMG_FILL_LAYER_OPACITY = 60

TOPIC_CLASS = "font-weight-bold"
ICON_WARN_CLASS = "fa-solid fa-triangle-exclamation fa-lg pr-1 text-warning"
P_WARN_CLASS = ""
P_WARN_STYLE = "border: 1px solid #ebebeb; border-radius: 0.25rem; padding: 0.5rem; background-color: rgba(241,99,52,.1);"

FTP_HOST = parser.get("FTP", "FTP_HOST")
FTP_USER = parser.get("FTP", "FTP_USER")
FTP_PASS = parser.get("FTP", "FTP_PASS")
FTP_IMG_PATH = parser.get("FTP", "FTP_IMG_PATH")

SITE = parser.get("SITE", "SITE")
SITE_IMAGES_PATH = parser.get("SITE", "SITE_IMAGES_PATH")

JOOMLA_ANCOR_IMAGE = "#joomlaImage://local-images/forecasts/day/"
JOOMLA_ANCOR_INTRO_IMG_SIZE = f"?width={INTRO_IMG_SIZE[0]}&height={INTRO_IMG_SIZE[1]}"
JOOMLA_ANCOR_FULL_IMG_SIZE = f"?width={FULL_IMG_SIZE[0]}&height={FULL_IMG_SIZE[1]}"

CREATED_BY_ID = 388

LANG = ""

MONTHS = {
    "ua": {
        "01": ["січень", "січня"],
        "02": ["лютий", "лютого"],
        "03": ["березень", "березня"],
        "04": ["квітень", "квітня"],
        "05": ["травень", "травня"],
        "06": ["червень", "червня"],
        "07": ["липень", "липня"],
        "08": ["серпень", "серпня"],
        "09": ["вересень", "вересня"],
        "10": ["жовтень", "жовтня"],
        "11": ["листопад", "листопада"],
        "12": ["грудень", "грудня"],
    },
    "ru": {
        "01": ["январь", "января"],
        "02": ["февраль", "февраля"],
        "03": ["март", "марта"],
        "04": ["апрель", "апреля"],
        "05": ["май", "мая"],
        "06": ["июнь", "июня"],
        "07": ["июль", "июля"],
        "08": ["август", "августа"],
        "09": ["сентябрь", "сентября"],
        "10": ["октябрь", "октября"],
        "11": ["ноябрь", "ноября"],
        "12": ["декабрь", "декабря"],
    },
}

SEPTENER = {
    "ua": {
        "0": ["Місяць", "Місяця"],
        "1": ["Марс", "Марса"],
        "2": ["Меркурій", "Меркурія"],
        "3": ["Юпітер", "Юпітера"],
        "4": ["Венера", "Венери"],
        "5": ["Сатурн", "Сатурна"],
        "6": ["Сонце", "Сонця"],
    },
    "ru": {
        "0": ["Луна", "Луны"],
        "1": ["Марс", "Марса"],
        "2": ["Меркурий", "Меркурия"],
        "3": ["Юпитер", "Юпитера"],
        "4": ["Венера", "Венеры"],
        "5": ["Сатурн", "Сатурна"],
        "6": ["Солнце", "Солнца"],
    },
}

LABELS = {
    "ua": {
        "header_line_date": re.compile(r"^(\d{1,2}\.\d{1,2}\.\d{4})"),
        "header_line_moon_day": re.compile(
            r"([Мм]ісячна доба) (\d{1,2}/\d{1,2}|\d{1,2})"
        ),
        "header_line_sun_day": re.compile(r"([СсCc]онячний день) (\d{1,2})"),
        "header_line_day_manager": re.compile(r"([Дд]ень під управлінням) (\w+)"),
        "moon_day": re.compile(r"^(\d{1,2}) (.*)"),
        "moon_day_start": re.compile(r"^([Пп]очаток)[:.]?[ ]?(.*)"),
        "moon_day_end": re.compile(r"^([Зз]акінчення)[:.]?[ ]?(.*)"),
        "moon_day_symbol": re.compile(r"^([СсCc]имвол) дня[:.]?[ ]?(.*)"),
        "moon_day_slogan": re.compile(r"^([Дд]евіз)[:.]?[ ]?(.*)"),
        "moon_day_satan": re.compile(r"^([СсCc]атанинськ(?:ий день|а доба))"),
        "moon_day_favorable": re.compile(r"^([СсCc]приятливо)[:.]?[ ]?(.*)"),
        "moon_day_important": re.compile(r"^([Вв]ажливо)[:.]?[ ]?(.*)"),
        "moon_day_dreams": re.compile(r"^([СсCc]ни)[:.]?[ ]?(.*)"),
        "moon_in_zodiac_sign": re.compile(r"^([Мм]ісяць [ув] .*)"),
        "moon_in_zodiac_sign_helth": re.compile(r"^([Зз]доров.я)[:.]?[ ]?(.*)"),
        "moon_in_zodiac_sign_void_moon": re.compile(
            r"^([Мм]ісяць без курсу)[:.]?[ ]?(.*)"
        ),
        "day_aspects": re.compile(r"^([Аа]спекти дня)[:.]?[ ]?(.*)"),
        "day_forecast": re.compile(r"^([Пп]рогноз)[:.]?[ ]?(.*)"),
        "footer_remark": re.compile(r"^.*[Рр]екомендації мають загальний .*"),
        "footer_by_line": re.compile(r"^([Зз] любов.ю[,]?)"),
        "intro_ru": re.compile(
            r"^\[\d+\]\d{1,2}\.\d{1,2}\.\d{4} [Лл]унные сутки \d{1,2} [СсCc]олнечный день \d{1,2}$"
        ),
        "intro_ua": re.compile(
            r"^\[\d+\]\d{1,2}\.\d{1,2}\.\d{4} [Мм]ісячна доба \d{1,2} [СсCc]онячний день \d{1,2}$"
        ),
    },
    "ru": {
        "header_line_date": re.compile(r"^(\d{1,2}\.\d{1,2}\.\d{4})"),
        "header_line_moon_day": re.compile(
            r"([Лл]унные сутки) (\d{1,2}/\d{1,2}|\d{1,2})"
        ),
        "header_line_sun_day": re.compile(r"([СсCc]олнечный день) (\d{1,2})"),
        "header_line_day_manager": re.compile(r"([Дд]ень под управлением) (\w+)"),
        "moon_day": re.compile(r"^(\d{1,2}) (.*)"),
        "moon_day_start": re.compile(r"^([Нн]ачало)[:.]?[ ]?(.*)"),
        "moon_day_end": re.compile(r"^([Оо]кончание)[:.]?[ ]?(.*)"),
        "moon_day_symbol": re.compile(r"^([СсCc]имвол) дня[:.]?[ ]?(.*)"),
        "moon_day_slogan": re.compile(r"^([Дд]евиз)[:.]?[ ]?(.*)"),
        "moon_day_satan": re.compile(r"^([СсCc]атанинск(?:ий день|ие сутки))"),
        "moon_day_favorable": re.compile(r"^([Бб]лагоприятно)[:.]?[ ]?(.*)"),
        "moon_day_important": re.compile(r"^([Вв]ажно)[:.]?[ ]?(.*)"),
        "moon_day_dreams": re.compile(r"^([СсCc]ны)[:.]?[ ]?(.*)"),
        "moon_in_zodiac_sign": re.compile(r"^([Лл]уна (?:в|во) .*)"),
        "moon_in_zodiac_sign_helth": re.compile(r"^([Зз]доровье)[:.]?[ ]?(.*)"),
        "moon_in_zodiac_sign_void_moon": re.compile(
            r"^([Лл]уна без курса)[:.]?[ ]?(.*)"
        ),
        "day_aspects": re.compile(r"^([Аа]спекты дня)[:.]?[ ]?(.*)"),
        "day_forecast": re.compile(r"^([Пп]рогноз)[:.]?[ ]?(.*)"),
        "footer_remark": re.compile(r"^.*[Рр]екомендации носят общий .*"),
        "footer_by_line": re.compile(r"^([СсCc] любовью[,]?)"),
        "intro_ru": re.compile(
            r"^\[\d+\]\d{1,2}\.\d{1,2}\.\d{4} [Лл]унные сутки \d{1,2} [СсCc]олнечный день \d{1,2}$"
        ),
        "intro_ua": re.compile(
            r"^\[\d+\]\d{1,2}\.\d{1,2}\.\d{4} [Мм]ісячна доба \d{1,2} [СсCc]онячний день \d{1,2}$"
        ),
    },
}

LANGUAGE_OVERRIDES = {
    "ua": {
        "HTML_CONTENT_A_P1": "Детальний астрологічний прогноз на день",
        "HTML_CONTENT_ID_P1": "Детальний астрологічний прогноз на день",
        "HTML_CONTENT_A_P2": "місячна доба: на що слід звернути увагу",
        "HTML_CONTENT_ID_P2": "місячна доба: основні характеристики",
        "HTML_CONTENT_A_P3": "Положення Місяця в Зодіаку та його вплив",
        "HTML_CONTENT_ID_P3": "Положення Місяця в Зодіаку та його вплив",
        "HTML_CONTENT_SATAN_DAY": "Сатанинська місячна доба",
        "HTML_CONTENT_MOON_POSITION": "Положення Місяця",
        "HTML_CONTENT_NATAL_TEXT": "індивідуальна консультація",
        "HTML_CONTENT_NATAL_LINK": "index.php?option=com_content&amp;view=article&amp;id=65:individualnyj-goroskop-natalnaya-karta&amp;catid=28:consultations&amp;lang=uk-UA&amp;Itemid=354",
        "HTML_CONTENT_VOID_MOON_TEXT": r"календар Місяц.? без курсу",
        "HTML_CONTENT_VOID_MOON_LINK": "index.php?option=com_content&amp;view=article&amp;id=1562:void-moon-2024&amp;catid=34:calendars&amp;lang=uk-UA&amp;Itemid=398",
        "HTML_CONTENT_FOOTER": "З любов’ю, Астролог Людмила Давиденко",
        "W_YEAR": ["рік", "року"],
        "W_PICTURE": ["картинка"],
        "SITE_CATEGORY_ID": 33,
        "SITE_LANGUAGE_CODE": "uk-UA",
        "SITE_ASTROLOGICAL_FORECAST_FOR": "Астрологічний прогноз на",
        "SITE_ASTROLOGER_FORECAST_FOR": "Прогноз астролога на",
    },
    "ru": {
        "HTML_CONTENT_A_P1": "Детальный астрологический прогноз на день",
        "HTML_CONTENT_ID_P1": "Детальный астрологический прогноз на день",
        "HTML_CONTENT_A_P2": "лунные сутки: на что следует обратить внимание",
        "HTML_CONTENT_ID_P2": "лунные сутки: основные характеристики",
        "HTML_CONTENT_A_P3": "Положение Луны в Зодиаке и ее влияние",
        "HTML_CONTENT_ID_P3": "Положение Луны в Зодиаке и ее влияние",
        "HTML_CONTENT_SATAN_DAY": "Сатанинские лунные сутки",
        "HTML_CONTENT_MOON_POSITION": "Положение Луны",
        "HTML_CONTENT_NATAL_TEXT": "индивидуальная консультация",
        "HTML_CONTENT_NATAL_LINK": "index.php?option=com_content&amp;view=article&amp;id=1604:individualnyj-goroskop-natalnaya-karta&amp;catid=50:consultations-ru&amp;lang=ru-RU&amp;Itemid=614",
        "HTML_CONTENT_VOID_MOON_TEXT": r"календарь [Лл]ун.? без курса",
        "HTML_CONTENT_VOID_MOON_LINK": "index.php?option=com_content&amp;view=article&amp;id=1563:void-moon-2024&amp;catid=68:calendars-ru&amp;lang=ru-RU&amp;Itemid=628",
        "HTML_CONTENT_FOOTER": "С любовью, Астролог Людмила Давиденко",
        "W_YEAR": ["год", "года"],
        "W_PICTURE": ["картинка"],
        "SITE_CATEGORY_ID": 57,
        "SITE_LANGUAGE_CODE": "ru-RU",
        "SITE_ASTROLOGICAL_FORECAST_FOR": "Астрологический прогноз на",
        "SITE_ASTROLOGER_FORECAST_FOR": "Прогноз астролога на",
    },
}


def log(log_str: str, severity="WARNING") -> None:
    print(f"\n-> {severity} - Key '{log_str}' not found!\n")


def dict_check_key(data: dict, key: str):

    if data.get(key, None):

        return True

    else:
        log(key)

        return False


def str_first_letter_to_upper(src_str: str) -> str:

    if not src_str:
        return ""

    return src_str[0].upper() + src_str[1:]


def str_first_letter_to_lower(src_str: str) -> str:

    if not src_str:
        return ""

    return src_str[0].lower() + src_str[1:]


def str_normalize(src_str: str) -> str:

    result = src_str.strip().strip("-– ")

    if result.startswith("#") or result.startswith("http"):
        result = ""

    if result != "":

        result = re.sub(r"\s{2,}", " ", result)

        result = re.sub(r"([(«/]) ", r"\1", result)

        result = re.sub(r"([а-яА-Я]+)([(«])([а-яА-Я]+)", r"\1 \2\3", result)

        result = re.sub(r" ([)»,.:/])", r"\1", result)
        result = re.sub(r"([а-яА-Я]+)([)».,:])([а-яА-Я]+)", r"\1\2 \3", result)

        result = re.sub(r"([а-яА-Я]+)(–)([а-яА-Я]+)", r"\1 \2 \3", result)
        result = result.replace(" - ", " – ")

    return result


def str_default_normalize(src_str: str) -> str:

    allow_char = "().,:;/"

    result = src_str.strip().strip("-– ")

    if result != "":

        result = re.sub(r"\s{2,}", " ", result)

        result = re.sub(r"\s*/\s*", "/", result)

        result = re.sub(r"\s*([(«])\s*", r" \1", result)
        result = re.sub(r"(\w+)([(«])(\w+)", r"\1 \2\3", result)

        result = re.sub(r"\s*([)»,.:])", r"\1", result)
        result = re.sub(r"(\w+)([)».,:])(\w+)", r"\1\2 \3", result)

        result = re.sub(r"(\w+)(–)(\w+)", r"\1 \2 \3", result)

        result = re.sub(r"(\w+)- ", r"\1 – ", result)
        result = re.sub(r" -(\w+)", r" – \1", result)

        result = result.replace(" - ", " – ")

    return result


def join_data(*data, sep=" ") -> str:

    result = ""

    if not data:
        return result

    for item in filter(None, data):
        if type(item) == list:
            item = clean_list(item)
            if item:
                result = result + sep + sep.join(item) if result else sep.join(item)
        elif type(item) == str:
            result = result + sep + item if result else item
        else:
            continue

    return result


def clean_list(data: list) -> list:

    result = []

    if not data:
        return result

    for elem in filter(None, data):
        if elem.strip():
            result.append(elem.strip())
        else:
            continue

    return result


def prepare_list_content(data: list, sep=":", end=".", first_letter="") -> list:

    result = []

    data = clean_list(data)

    if not data:
        return result

    first_item = data[0].rstrip(sep)
    result.append(str_first_letter_to_upper(first_item) + sep)

    for item in data[1:]:

        if end == "." and re.findall(r"(?:[?!]|\.\.\.)$", item):
            cur_end = ""
        else:
            item = item.rstrip(end)
            cur_end = end

        if first_letter == "upper":
            result.append(str_first_letter_to_upper(item) + cur_end)
        elif first_letter == "lower":
            result.append(str_first_letter_to_lower(item) + cur_end)
        else:
            result.append(item + cur_end)

    return result


def find_label_positions(text_list: list) -> dict:

    label_positions = {}

    if text_list == []:
        return label_positions

    for idx, item in enumerate(text_list):
        for key in LABELS[LANG].keys():
            if key in label_positions:
                continue
            else:
                result = re.findall(LABELS[LANG][key], item)
                if result:
                    label_positions[key] = idx

    return label_positions


def get_data_from_text(text_list: list) -> dict:

    result_data = {}

    if text_list == []:
        return result_data

    label_positions = find_label_positions(text_list)

    text_points = list(set(label_positions.values()))
    text_points.sort()

    for label, label_pos in label_positions.items():

        label_idx = text_points.index(label_pos)

        if text_points[label_idx] != text_points[-1]:
            next_label_pos = text_points[label_idx + 1]
        else:
            next_label_pos = len(text_list)

        block_length = next_label_pos - label_pos

        find_result = re.findall(LABELS[LANG][label], text_list[label_pos])[0]

        if type(find_result) == str:
            result_data.setdefault(label, []).append(find_result)
        if type(find_result) == tuple:
            for elem in filter(None, find_result):
                result_data.setdefault(label, []).append(elem)

        i = 1

        while i < block_length:
            result_data[label].append(text_list[label_pos + i])
            i += 1

    return result_data


def html_tag(*data, tag="p", _id="", _class="", _style="", a_href="") -> str:

    result = ""

    eol = "" if tag in ("a", "i", "span") else "\n"

    tag_body = join_data(*data) if tag != "ul" else join_data(*data, sep="</li>\n<li>")

    tag_open = f"<{tag}"
    tag_open = f'{tag_open} id="{_id}"' if _id else tag_open
    tag_open = f'{tag_open} class="{_class}"' if _class else tag_open
    tag_open = f'{tag_open} style="{_style}"' if _style else tag_open
    tag_open = f'{tag_open} href="{a_href}"' if a_href and tag == "a" else tag_open
    tag_open = f"{tag_open}>" if tag != "ul" else f"{tag_open}>\n<li>"

    if tag == "ul":
        tag_close = "</li>\n</ul>"
    elif tag in ("hr",):
        tag_close = ""
    else:
        tag_close = f"</{tag}>"

    result = (
        join_data(tag_open, tag_body, tag_close, eol, sep="")
        if tag_body or tag in ("a", "i", "hr")
        else ""
    )

    return result


def html_formatted_block(data: list, format="p") -> str:
    """
    data = ["elem0", "elem1", "elem2", ..., "elemN"], N >= 1

    elem0 - header, may be "",
    elem1, ..., elemN - elements, if "", than skip

    1. format = ul:
        <p><span>elem0</span></p>
        <ul>
        <li>elem1</li>
        <li>elem2</li>
        ...
        <li>elemN</li>
        </ul>

    2. format = p-list:
        <p><span>elem0</span></p>
        <p>elem1</p>
        <p>elem2</p>
        ...
        <p>elemN</p>

    3. format = p-list-ext:
        <p><span>elem0</span>elem1</p>
        <p>elem2</p>
        ...
        <p>elemN</p>

    4. format = p:
        <p><span>elem0</span>elem1 elem2 ... elemN</p>
    """

    result = ""

    data = clean_list(data)

    if not data or len(data) == 1:
        return result

    span_block = html_tag(data[0], tag="span", _class=TOPIC_CLASS)
    p_block = html_tag(data[0], tag="p", _class=TOPIC_CLASS)

    data = data[1:]

    if format == "ul":

        result = p_block + html_tag(data, tag="ul")

    elif format == "p-list":

        result = p_block + "".join(map(lambda x: html_tag(x, tag="p"), data))

    elif format == "p-list-ext":

        result = html_tag(span_block, data[0], tag="p") + "".join(
            map(lambda x: html_tag(x, tag="p"), data[1:])
        )

    else:

        result = html_tag(span_block, " ".join(data), tag="p")

    return result


def get_forecast_date(data: dict, form="num", sep="-", start="day"):

    # form: num, dt, text

    if data.get("header_line_date", None):

        try:
            dt = datetime.strptime(data["header_line_date"][0], "%d.%m.%Y").date()
        except:
            print("CRITICAL ERROR IN get_forecast_date")
            exit()

        if form == "text":
            day, month, year = dt.strftime("%d.%m.%Y").split(".")

            day = day.removeprefix("0")
            month = MONTHS[LANG][month][1]

            return f"{day} {month} {year}"
        elif form == "dt":
            return dt

        if start == "day":
            return dt.strftime(f"%d{sep}%m{sep}%Y")
        else:
            return dt.strftime(f"%Y{sep}%m{sep}%d")

    else:
        log("header_line_date", "CRITICAL")
        exit()


def make_html_block(data: dict, key: str, params: list) -> str:

    result = ""

    if not data or not key:
        return result

    format, sep, end, fl = params

    if dict_check_key(data, key):

        if key == "moon_in_zodiac_sign_helth":
            for i in range(0, len(data[key])):
                data[key][i] = data[key][i].replace(":", "")

        content = prepare_list_content(data[key], sep=sep, end=end, first_letter=fl)

        result = html_formatted_block(content, format=format)

    return result


def make_html_link(src_str: str, text: str, link: str) -> str:

    result = ""

    link_tag_open = '<a href="' + link + '">'
    link_tag_close = "</a>"

    result = re.sub(f"({text})", link_tag_open + r"\1" + link_tag_close, src_str)

    return result


def make_html_warning_str(src_str: str) -> str:

    return join_data(html_tag(tag="i", _class=ICON_WARN_CLASS), src_str)


def make_html_content(data: dict) -> str:

    out = ""

    #######################

    if dict_check_key(data, "intro_" + LANG):
        items = data["intro_" + LANG][1].split(": ")
        data["intro_" + LANG][1] = str_first_letter_to_upper(items[1])

        for item in data["intro_" + LANG][1:]:
            out += html_tag(item, tag="p")

    out += html_tag(tag="hr", _id="system-readmore")

    #######################

    header_forecast_date = get_forecast_date(data, form="text")

    header_line_day_manager = (
        " ".join(data["header_line_day_manager"])
        if dict_check_key(data, "header_line_day_manager")
        else "______"
    )

    header_line_moon_day = (
        " ".join(data["header_line_moon_day"][::-1])
        if dict_check_key(data, "header_line_moon_day")
        else "______"
    )

    header_line_sun_day = (
        " ".join(data["header_line_sun_day"][::-1])
        if dict_check_key(data, "header_line_sun_day")
        else "______"
    )

    moon_day = (
        data["moon_day"][0].removeprefix("0")
        if dict_check_key(data, "moon_day")
        else "______"
    )

    out += html_tag(
        header_forecast_date, "-", header_line_day_manager, tag="p", _class=TOPIC_CLASS
    )
    out += html_tag(
        header_line_moon_day + ",",
        header_line_sun_day + ":",
        tag="p",
        _class=TOPIC_CLASS,
    )

    out += html_tag(
        html_tag(LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_A_P1"], tag="a", a_href="#p1"),
        html_tag(
            moon_day,
            LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_A_P2"],
            tag="a",
            a_href="#p2",
        ),
        html_tag(LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_A_P3"], tag="a", a_href="#p3"),
        tag="ul",
    )

    #######################

    out += html_tag(html_tag(tag="a", _id="p1"), tag="p")
    out += html_tag(LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_ID_P1"], tag="h2")

    out += make_html_block(data, "day_aspects", ["ul", ":", ".", "upper"])

    if dict_check_key(data, "day_forecast"):
        lead_text = f'{data["day_forecast"][0]} на {header_forecast_date}'
        out += html_formatted_block(
            prepare_list_content([lead_text, *data["day_forecast"][1:]]),
            format="p-list",
        )

    if dict_check_key(data, "footer_remark"):
        remark = make_html_link(
            data["footer_remark"][0],
            LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_NATAL_TEXT"],
            LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_NATAL_LINK"],
        )

        out += html_tag(
            make_html_warning_str(remark),
            tag="p",
            _class=P_WARN_CLASS,
            _style=P_WARN_STYLE,
        )

    #######################

    out += html_tag(html_tag(tag="a", _id="p2"), tag="p")
    out += html_tag(moon_day, LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_ID_P2"], tag="h2")

    out += make_html_block(data, "moon_day_start", ["p", ":", ".", "lower"])
    out += make_html_block(data, "moon_day_end", ["p", ":", ".", "lower"])
    out += make_html_block(data, "moon_day_symbol", ["p", ":", ".", "lower"])
    out += make_html_block(data, "moon_day_slogan", ["p", ":", ".", "lower"])

    if moon_day in ("9", "15", "19", "23", "29"):
        out += html_tag(
            prepare_list_content(
                [LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_SATAN_DAY"]], sep="."
            ),
            tag="p",
            _class=TOPIC_CLASS,
        )

    out += make_html_block(data, "moon_day_favorable", ["ul", ":", ";", "lower"])
    out += make_html_block(data, "moon_day_important", ["ul", ":", ";", "lower"])
    out += make_html_block(data, "moon_day_dreams", ["p", ".", ".", "upper"])

    #######################

    out += html_tag(html_tag(tag="a", _id="p3"), tag="p")
    out += html_tag(LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_ID_P3"], tag="h2")

    if dict_check_key(data, "moon_in_zodiac_sign"):
        out += html_formatted_block(
            prepare_list_content(
                [
                    LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_MOON_POSITION"],
                    *data["moon_in_zodiac_sign"],
                ]
            )
        )

    out += make_html_block(data, "moon_in_zodiac_sign_helth", ["p", ".", ".", "upper"])

    if dict_check_key(data, "moon_in_zodiac_sign_void_moon"):

        item_idx_1 = data["moon_in_zodiac_sign_void_moon"][1]

        match = re.search(
            LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_VOID_MOON_TEXT"], item_idx_1
        )
        pos_1 = match.start()
        pos_2 = item_idx_1.rfind(".", 0, pos_1) + 1

        data["moon_in_zodiac_sign_void_moon"][1] = item_idx_1[0:pos_2:]
        data["moon_in_zodiac_sign_void_moon"].insert(2, item_idx_1[pos_2:])

        data["moon_in_zodiac_sign_void_moon"][2] = make_html_link(
            data["moon_in_zodiac_sign_void_moon"][2],
            LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_VOID_MOON_TEXT"],
            LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_VOID_MOON_LINK"],
        )

        void_moon_data = prepare_list_content(data["moon_in_zodiac_sign_void_moon"])

        out += html_formatted_block(void_moon_data[:2], format="p")

        out += html_tag(
            make_html_warning_str(void_moon_data[2]),
            tag="p",
            _class=P_WARN_CLASS,
            _style=P_WARN_STYLE,
        )

        for item in void_moon_data[3:]:
            out += html_tag(item, tag="p")

    #######################

    out += html_tag(
        LANGUAGE_OVERRIDES[LANG]["HTML_CONTENT_FOOTER"], tag="p", _class="float-right"
    )

    out = out.strip()

    return out


def upload_to_file(data: str) -> None:

    data = (
        '<link href="https://stackpath.bootstrapcdn.com/bootstrap/4.4.1/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-Vkoo8x4CGsO3+Hhxv8T/Q5PaXtkKtu6ug5TOeNV6gBiFeWPGFN9MuhOf23Q9Ifjh" crossorigin="anonymous">\n'
        + data
    )
    data = (
        '<link href="https://zodiac.kiev.ua/templates/masterbootstrap/css/template.css" rel="stylesheet">\n'
        + data
    )
    data = (
        '<link href="https://zodiac.kiev.ua/templates/masterbootstrap/css/template.custom.css" rel="stylesheet"><div style="padding: 20px;">\n'
        + data
    )

    with open("src/index.html", "w", encoding="utf8") as fh:
        fh.write(data)


def upload_to_site(data) -> None:
    url = parser.get("JOOMLA", "API_URL")

    payload = json.dumps(data)
    headers = {
        "Content-Type": "application/json",
        "Authorization": parser.get("JOOMLA", "API_KEY"),
    }

    response = requests.request("POST", url, headers=headers, data=payload)
    print(response.text)


def forecast_upload(text_list: list, do_upload_to_site: str) -> None:

    data = get_data_from_text(text_list)

    html_output = make_html_content(data)

    upload_to_file(html_output)

    if do_upload_to_site == "y":

        delta_1_day = timedelta(days=1)
        delta_2_day = timedelta(days=2)

        forecast_date = get_forecast_date(data, form="text")
        forecast_date_dt = get_forecast_date(data, form="dt")
        forecast_date_y_m_d = get_forecast_date(data, start="year")
        forecast_date_d_m_y = get_forecast_date(data)

        intro_text = html_output.split('<hr id="system-readmore">')[0]
        intro_text = intro_text.replace("<p>", "").replace("</p>", "")
        metadesc = join_data(
            LANGUAGE_OVERRIDES[LANG]["SITE_ASTROLOGER_FORECAST_FOR"],
            forecast_date + ".",
            intro_text,
        )[:159:]

        article = dict()
        article["images"] = dict()
        article["metadata"] = dict()

        article.update(
            {
                "title": join_data(
                    LANGUAGE_OVERRIDES[LANG]["SITE_ASTROLOGICAL_FORECAST_FOR"],
                    forecast_date,
                )
            }
        )
        article.update({"alias": forecast_date_d_m_y})
        article.update({"articletext": html_output})
        article.update({"catid": LANGUAGE_OVERRIDES[LANG]["SITE_CATEGORY_ID"]})
        # article.update({"featured": 0})
        article.update({"language": LANGUAGE_OVERRIDES[LANG]["SITE_LANGUAGE_CODE"]})
        # article.update({"tags": []})
        # article.update({"note": ""})

        article["images"].update(
            {
                "image_intro": join_data(
                    SITE_IMAGES_PATH,
                    forecast_date_y_m_d,
                    INTRO_IMG_TAIL,
                    JOOMLA_ANCOR_IMAGE,
                    forecast_date_y_m_d,
                    INTRO_IMG_TAIL,
                    JOOMLA_ANCOR_INTRO_IMG_SIZE,
                    sep="",
                )
            }
        )
        article["images"].update(
            {
                "image_intro_alt": join_data(
                    LANGUAGE_OVERRIDES[LANG]["SITE_ASTROLOGICAL_FORECAST_FOR"],
                    forecast_date,
                    LANGUAGE_OVERRIDES[LANG]["W_YEAR"][1],
                    LANGUAGE_OVERRIDES[LANG]["W_PICTURE"][0],
                )
            }
        )
        article["images"].update({"float_intro": ""})
        article["images"].update(
            {
                "image_intro_caption": join_data(
                    LANGUAGE_OVERRIDES[LANG]["SITE_ASTROLOGICAL_FORECAST_FOR"],
                    forecast_date,
                    LANGUAGE_OVERRIDES[LANG]["W_YEAR"][1],
                )
            }
        )

        article["images"].update(
            {
                "image_fulltext": join_data(
                    SITE_IMAGES_PATH,
                    forecast_date_y_m_d,
                    FULL_IMG_TAIL,
                    JOOMLA_ANCOR_IMAGE,
                    forecast_date_y_m_d,
                    FULL_IMG_TAIL,
                    JOOMLA_ANCOR_FULL_IMG_SIZE,
                    sep="",
                )
            }
        )
        article["images"].update(
            {
                "image_fulltext_alt": join_data(
                    LANGUAGE_OVERRIDES[LANG]["SITE_ASTROLOGICAL_FORECAST_FOR"],
                    forecast_date,
                    LANGUAGE_OVERRIDES[LANG]["W_YEAR"][1],
                    LANGUAGE_OVERRIDES[LANG]["W_PICTURE"][0],
                )
            }
        )
        article["images"].update({"float_fulltext": "none"})
        article["images"].update(
            {
                "image_fulltext_caption": join_data(
                    LANGUAGE_OVERRIDES[LANG]["SITE_ASTROLOGICAL_FORECAST_FOR"],
                    forecast_date,
                    LANGUAGE_OVERRIDES[LANG]["W_YEAR"][1],
                )
            }
        )

        article.update(
            {"publish_up": str(forecast_date_dt - delta_1_day) + " 00:00:01"}
        )
        # article.update({"publish_down": None})
        # article.update({"featured_up": None})
        # article.update({"featured_down": None})
        article.update({"created": str(forecast_date_dt - delta_2_day) + " 09:00:00"})
        article.update({"created_by": CREATED_BY_ID})
        article.update({"created_by_alias": ""})

        article.update({"metadesc": metadesc})
        article.update({"metakey": ""})
        article["metadata"].update({"robots": "index, follow"})
        article["metadata"].update({"author": ""})
        article["metadata"].update({"rights": ""})

        print(article)

        upload_to_site(article)

        print(
            f"""
    ************************************************************************************
              
        Перед публікацією статті необхідно виконати наступне:
              
        1. Редактирование материала -> Отображение -> Заголовок
            Вибрати "Скрыть"
              
        2. Редактирование материала -> Отображение -> Заголовок страницы в браузере
            Вставити | {LANGUAGE_OVERRIDES[LANG]["SITE_ASTROLOGER_FORECAST_FOR"]} {forecast_date} {LANGUAGE_OVERRIDES[LANG]["W_YEAR"][1]} |

        3. Редактирование материала -> Публикация -> Метаданные -> Метатег Description        
            Перевірити і виправити при необхідності

        4. Редактирование материала -> Связи
            Налаштувати відповідність UA-RU (якщо обидві статті вже є на сайті)

        Після публікації статті не забути налаштувати TAGZ Open Graph для неї.

    ************************************************************************************
              """
        )


def forecast_read(forecast_date: str) -> list:

    forecast_date_year, forecast_date_month, _ = forecast_date.split("-")

    forecast_file_name = f"{forecast_date}_{LANG}.docx"
    forecast_intro_file_name = f"{forecast_date}_intro.docx"

    forecast_file_path = Path(
        FORECAST_ROOT_FOLDER,
        forecast_date_year,
        forecast_date_month,
        forecast_file_name,
    )

    forecast_intro_file_path = Path(
        FORECAST_ROOT_FOLDER,
        forecast_date_year,
        forecast_date_month,
        forecast_intro_file_name,
    )

    if not forecast_file_path.exists() and not forecast_intro_file_path.exists():
        print(f"-> Не знайдено файл з прогнозом {forecast_file_name}!")
        print(f"-> Не знайдено файл з коротким описом {forecast_intro_file_name}!")
        return []

    elif not forecast_file_path.exists():
        print(f"-> Не знайдено файл з прогнозом {forecast_file_name}!")
        return []

    elif not forecast_intro_file_path.exists():
        print(f"-> Не знайдено файл з коротким описом {forecast_intro_file_name}!")
        return []

    document_forecast = Document(forecast_file_path)
    document_intro = Document(forecast_intro_file_path)

    text = []

    for number, file in enumerate((document_forecast, document_intro)):
        for paragraph in file.paragraphs:
            paragraph_text = str_normalize(paragraph.text)
            if paragraph_text:
                if number == 0:
                    text.append(paragraph_text)
                else:
                    text.append(f"[{number}]{paragraph_text}")

    return text


def make_forecast_img(forecast_img_date: str) -> bool:

    forecast_img_date_year, forecast_img_date_month, forecast_img_date_day = (
        forecast_img_date.split("-")
    )

    forecast_img_src_file_name = f'{forecast_img_date_day.lstrip("0")}.JPG'

    forecast_img_src_file_path = Path(
        FORECAST_ROOT_FOLDER,
        forecast_img_date_year,
        forecast_img_date_month,
        FORECAST_IMG_SRC_FOLDER,
        forecast_img_src_file_name,
    )

    forecast_img_dst_files_path = Path(
        FORECAST_ROOT_FOLDER,
        forecast_img_date_year,
        forecast_img_date_month,
        FORECAST_IMG_DST_FOLDER,
    )

    if not forecast_img_src_file_path.exists():
        print(
            f"-> Не знайдено файл зображення для прогнозу {forecast_img_src_file_name}!"
        )
        return False

    with Image.open(forecast_img_src_file_path) as img:
        img.load()

    ratio_to_full = img.height / FULL_IMG_SIZE[1]
    tmp_img_width = int(img.width / ratio_to_full)

    tmp_img = img.resize((tmp_img_width, FULL_IMG_SIZE[1]))
    tmp_img_half_width = tmp_img.width // 2

    full_img = Image.new("RGB", FULL_IMG_SIZE, (255, 255, 255))
    full_img_half_width = full_img.width // 2

    full_img.paste(tmp_img, (full_img_half_width - tmp_img_half_width, 0))

    ###################

    ratio_to_fb = full_img.height / FB_IMG_SIZE[1]
    tmp_img_width = int(full_img.width / ratio_to_fb)

    tmp_img = full_img.resize((tmp_img_width, FB_IMG_SIZE[1]))

    crop_len = (tmp_img.width - FB_IMG_SIZE[0]) // 2

    fb_img = tmp_img.crop((crop_len, 0, crop_len + FB_IMG_SIZE[0], FB_IMG_SIZE[1]))

    fb_img.save(Path(forecast_img_dst_files_path, forecast_img_date + FB_IMG_TAIL))

    ###################

    ratio_to_intro = full_img.height / INTRO_IMG_SIZE[1]
    tmp_img_width = int(full_img.width / ratio_to_intro)

    tmp_img = full_img.resize((tmp_img_width, INTRO_IMG_SIZE[1]))

    crop_len = (tmp_img.width - INTRO_IMG_SIZE[0]) // 2

    intro_img = tmp_img.crop(
        (crop_len, 0, crop_len + INTRO_IMG_SIZE[0], INTRO_IMG_SIZE[1])
    )

    intro_img.save(
        Path(forecast_img_dst_files_path, forecast_img_date + INTRO_IMG_TAIL)
    )

    ###################

    transparent_layer = Image.new("RGB", FULL_IMG_SIZE, FULL_IMG_FILL_LAYER_COLOR)
    transparent_layer.putalpha(int(255 * FULL_IMG_FILL_LAYER_OPACITY / 100))

    full_img.paste(transparent_layer, (0, 0), transparent_layer)

    full_img.save(Path(forecast_img_dst_files_path, forecast_img_date + FULL_IMG_TAIL))

    print(f"-> Зображення успішно згенеровані!")
    return True


def upload_forecast_img(forecast_img_date: str) -> bool:

    forecast_img_date_year, forecast_img_date_month, _ = forecast_img_date.split("-")

    forecast_img_dst_files_path = Path(
        FORECAST_ROOT_FOLDER,
        forecast_img_date_year,
        forecast_img_date_month,
        FORECAST_IMG_DST_FOLDER,
    )

    full_forecast_img_filename = forecast_img_date + FULL_IMG_TAIL
    intro_forecast_img_filename = forecast_img_date + INTRO_IMG_TAIL
    fb_forecast_img_filename = forecast_img_date + FB_IMG_TAIL

    with ftputil.FTPHost(FTP_HOST, FTP_USER, FTP_PASS) as ftp_host:
        ftp_host.upload(
            Path(forecast_img_dst_files_path, full_forecast_img_filename),
            PurePosixPath(FTP_IMG_PATH, full_forecast_img_filename),
        )

        ftp_host.upload(
            Path(forecast_img_dst_files_path, intro_forecast_img_filename),
            PurePosixPath(FTP_IMG_PATH, intro_forecast_img_filename),
        )

        ftp_host.upload(
            Path(forecast_img_dst_files_path, fb_forecast_img_filename),
            PurePosixPath(FTP_IMG_PATH, fb_forecast_img_filename),
        )

    print(f"-> Зображення успішно завантажені на FTP сайту!")
    return True


def check_input(input_str: str, input_list: list) -> str:

    user_input = input(input_str)

    while user_input not in input_list:
        user_input = input(input_str)

    return user_input


def check_input_date(input_str: str) -> str:

    while True:
        user_input = input(input_str)
        if user_input == "":
            return ""

        result = re.fullmatch(
            r"^\d{4}-(?:0[1-9]|1[0-2])-(?:0[1-9]|[1-2][0-9]|3[0-1])$", user_input
        )

        if result:
            return result.group()


def check_file_exist(file: Path):
    pass


def main():

    this_date = str(datetime.now().date())

    global LANG

    while True:

        os.system("cls")

        print("\n*** Головне меню ***\n")
        print("1. Прогноз на день")
        print("2. Робота з зображеннями")
        print("3. Нормалізація тексту")
        print("\n0. Вихід")
        print("\n")

        action = (
            check_input("Зробіть свій вибір [1]: ", ["1", "2", "3", "0", ""]) or "1"
        )

        if action == "0":
            print("\nПрощавайте!\n")
            exit()

        elif action == "1":

            os.system("cls")

            print("\n*** Створення прогнозу ***\n")

            LANG = (
                check_input("<- Мова прогнозу (ua / ru) [ua]: ", ["ua", "ru", ""])
                or "ua"
            )

            forecast_date = (
                check_input_date(
                    f"<- Дата прогнозу в форматі РРРР-ММ-ДД [{this_date}]: "
                )
                or this_date
            )

            text = forecast_read(forecast_date)

            if text:

                forecast_upload_to_site = (
                    check_input(
                        "<- Публікувати прогноз на сайті? (y / n) [n]: ", ["y", "n", ""]
                    )
                    or "n"
                )

                forecast_upload(text, forecast_upload_to_site)

                print(f"-> Всі дії виконані!")

        elif action == "2":

            os.system("cls")

            print("\n*** Генерація зображень для прогнозу ***\n")

            forecast_img_date = (
                check_input_date(
                    f"<- Дата прогнозу, для якого генеруються зображення, в форматі РРРР-ММ-ДД [{this_date}]: "
                )
                or this_date
            )

            if make_forecast_img(forecast_img_date):

                forecast_img_upload_to_site_ftp = (
                    check_input(
                        "<- Завантажувати згенеровані зображення на FTP сайту? (y / n) [n]: ",
                        ["y", "n", ""],
                    )
                    or "n"
                )

                if forecast_img_upload_to_site_ftp == "y":

                    upload_forecast_img(forecast_img_date)

                print(f"-> Всі дії виконані!")

        elif action == "3":

            os.system("cls")

        input("<- Натисніть ENTER для продовження...")


if __name__ == "__main__":
    main()
